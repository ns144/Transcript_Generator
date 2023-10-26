from docx import Document
from docx.shared import Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from whisper.utils import format_timestamp
import uuid
from pathlib import Path


# Function to write the transcript in a DOCX file
def speakersToDoc(speaker_segments:list,translated_segments:list,scriptFilename:str, sourcefile="", translated=False):
    # DIN A4 page setup
    document = Document()
    document.sections[0].page_width = Mm(210)
    document.sections[0].page_height = Mm(297)

    # Document header
    document.add_heading('Transcription', 0)
    p = document.add_paragraph('File:  ' + Path(sourcefile).name)

    latest_timestamp = 0
    latest_index = 0

    # initialize table
    table = document.add_table(rows=1, cols=4)
    table.style = document.styles['Table Grid']
    hdr_cells = table.rows[0].cells

    paragraph = hdr_cells[0].paragraphs[0]
    run = paragraph.add_run('ID')
    run.bold = True
    hdr_cells[0].width = Mm(9.4)

    paragraph.alignment=WD_ALIGN_PARAGRAPH.CENTER
    paragraph = hdr_cells[1].paragraphs[0]
    run = paragraph.add_run('Timecode')
    run.bold = True
    hdr_cells[1].width = Mm(30)

    paragraph.alignment=WD_ALIGN_PARAGRAPH.CENTER
    paragraph = hdr_cells[2].paragraphs[0]
    run = paragraph.add_run('Text')
    hdr_cells[2].width = Mm(50)
    run.bold = True

    paragraph.alignment=WD_ALIGN_PARAGRAPH.CENTER
    paragraph = hdr_cells[3].paragraphs[0]
    run = paragraph.add_run('Translation')
    hdr_cells[3].width = Mm(50)
    run.bold = True
    paragraph.alignment=WD_ALIGN_PARAGRAPH.CENTER

    # add the speaker segments to the table
    for index_speaker,speaker in enumerate(speaker_segments):

        if len(speaker.segments) != 0:

            speaker_name = table.add_row()
            run = speaker_name.cells[0].paragraphs[0].add_run(speaker.speaker)
            speaker_name.cells[0].merge(speaker_name.cells[3])
            speaker_name.cells[0].paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
            run.bold = True


            trans_text = ""
            translation = translated_segments[index_speaker]
            start = 0
            end = 0

            for index, s in enumerate(speaker.segments):

                start = format_timestamp(s['start'],True,':')

                end = format_timestamp(s['end'],True,':')

                row_cells = table.add_row().cells

                row_cells[0].text = str(latest_index + index + 1)
                row_cells[0].width = Mm(9.4)
                row_cells[1].text = start + " - " + end
                row_cells[2].text = s['text']

                # When a translation is available add it into the table
                if translated:
                    trans_text = translation.segments[index]['text']

                row_cells[3].text = trans_text

                trans_text = ""

            latest_index += len(speaker.segments)

        

    paragraph = document.add_paragraph()
    run = paragraph.add_run('This transcript was generated automatically and needs further correction. Please ensure to have a translator check the content against the original audio.')
    run.italic = True   
    
    # When a file with the filename already exists, a uuid is added to the filename
    try:
        document.save(scriptFilename)
    except:
        path = Path(scriptFilename)
        scriptFilename = path.with_name(path.stem + '_' + uuid.uuid4().hex + '.docx')
        document.save(scriptFilename) #try again

    return scriptFilename


def transcriptToDoc(segments:list, translatedSegments:list, scriptFilename:str, sourcefile=""):

    #p.add_run('bold').bold = True
    #p.add_run(' and some ')
    #p.add_run('italic.').italic = True
    #document.add_page_break()
    #write table
    
    #read script
    document = Document()
    
    # A4 Size
    document.sections[0].page_width = Mm(210)
    document.sections[0].page_height = Mm(297)
    
    #Add header
    document.add_heading('ns144 Transcription', 0)
    p = document.add_paragraph('File:  ' + Path(sourcefile).name)
    
    table = document.add_table(rows=1, cols=4)
    table.style = document.styles['Table Grid']
    hdr_cells = table.rows[0].cells
    
    paragraph = hdr_cells[0].paragraphs[0]
    run = paragraph.add_run('ID')
    run.bold = True
    hdr_cells[0].width = Mm(9.4)
    paragraph.alignment=WD_ALIGN_PARAGRAPH.CENTER
    
    paragraph = hdr_cells[1].paragraphs[0]
    run = paragraph.add_run('Timecode')
    run.bold = True
    hdr_cells[1].width = Mm(30)
    paragraph.alignment=WD_ALIGN_PARAGRAPH.CENTER
    
    paragraph = hdr_cells[2].paragraphs[0]
    run = paragraph.add_run('Text')
    hdr_cells[2].width = Mm(50)
    run.bold = True
    paragraph.alignment=WD_ALIGN_PARAGRAPH.CENTER
    
    paragraph = hdr_cells[3].paragraphs[0]
    run = paragraph.add_run('Translation')
    hdr_cells[3].width = Mm(50)
    run.bold = True
    paragraph.alignment=WD_ALIGN_PARAGRAPH.CENTER
        
    trans_text = ""
    start = 0
    end = 0
    
    translation_size = len(translatedSegments)
    translationAvailable = translation_size > 0
    
    for index, segment in enumerate(segments):
        
        start = format_timestamp(segment['start'],True,':')
         
        #Add translated text
        if (translationAvailable and index <= translation_size-1):
            trans_text = translatedSegments[index]['text']
        
        end = format_timestamp(segment['end'],True,':')
        
        row_cells = table.add_row().cells
        
        row_cells[0].text = str(index + 1)
        row_cells[0].width = Mm(9.4)
        row_cells[1].text = start + " - " + end
        row_cells[2].text = segment['text']
        row_cells[3].text = trans_text

        trans_text = ""
    
    paragraph = document.add_paragraph()
    run = paragraph.add_run('This transcript was generated automatically and needs further correction. Please ensure to have a translator check the content against the original audio.')
    run.italic = True   
    
    try:
        document.save(scriptFilename)
    except:
        path = Path(scriptFilename)
        scriptFilename = path.with_name(path.stem + '_' + uuid.uuid4().hex + '.docx')
        document.save(scriptFilename) #try again

    return scriptFilename